"""从 PDF / DOCX / TXT / MD 中提取论文文本。"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

EXTRACT_SUFFIXES = {".pdf", ".docx", ".doc", ".txt", ".md"}


@dataclass
class PaperText:
    path: Path
    title: str = ""
    text: str = ""
    warnings: list[str] = field(default_factory=list)

    @property
    def n_chars(self) -> int:
        return len(self.text)


def _extract_pdf(path: Path) -> tuple[str, list[str]]:
    import fitz  # PyMuPDF

    warnings = []
    doc = fitz.open(path)
    parts = []
    for i, page in enumerate(doc):
        t = page.get_text("text").strip()
        if t:
            parts.append(t)
        else:
            warnings.append(f"第 {i + 1} 页未提取到文本（可能是扫描图片页）")
    doc.close()
    return "\n\n".join(parts), warnings


def _extract_docx(path: Path) -> tuple[str, list[str]]:
    import docx

    d = docx.Document(str(path))
    parts = []
    for para in d.paragraphs:
        t = para.text.strip()
        if t:
            prefix = ""
            if para.style is not None and para.style.name.startswith("Heading"):
                prefix = "## "
            parts.append(prefix + t)
    # 表格内容也纳入（数据、实验结果常在表格里）
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    warnings = []
    if not parts:
        warnings.append("DOCX 中未提取到文本")
    return "\n\n".join(parts), warnings


def _extract_text(path: Path) -> tuple[str, list[str]]:
    raw = path.read_bytes()
    for enc in ("utf-8", "gb18030", "utf-16"):
        try:
            return raw.decode(enc), []
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace"), ["文本编码无法识别，已按替换字符解码"]


def extract_paper(path: str | Path) -> PaperText:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix not in EXTRACT_SUFFIXES:
        raise ValueError(f"不支持的文件类型: {suffix}（支持 {sorted(EXTRACT_SUFFIXES)}）")

    if suffix == ".pdf":
        text, warnings = _extract_pdf(path)
    elif suffix in (".docx", ".doc"):
        if suffix == ".doc":
            raise ValueError("暂不支持旧版 .doc，请先另存为 .docx 或 PDF")
        text, warnings = _extract_docx(path)
    else:
        text, warnings = _extract_text(path)

    title = ""
    for line in text.splitlines():
        line = line.strip().lstrip("#").strip()
        if len(line) >= 4:
            title = line
            break

    return PaperText(path=path, title=title, text=text, warnings=warnings)


def collect_papers(folder_or_file: str | Path) -> list[Path]:
    """收集待批改的论文文件（自动跳过输出目录和隐藏文件）。"""
    p = Path(folder_or_file)
    if p.is_file():
        return [p]
    if p.is_dir():
        found = [
            f
            for f in sorted(p.rglob("*"))
            if f.is_file()
            and f.suffix.lower() in EXTRACT_SUFFIXES
            and not f.name.startswith(("~", "."))
            and "output" not in f.parts
        ]
        return found
    raise FileNotFoundError(p)
